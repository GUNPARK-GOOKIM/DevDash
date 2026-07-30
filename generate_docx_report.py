import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_comparison_docx():
    doc = Document()
    
    # Page setup - 0.5 in margins for 5-column landscape layout
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.0)
        section.page_height = Inches(8.5)
        
    # Document Header
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("DevDash vs. DBForge — 30-Parameter Engineering Comparison & Evaluation")
    r_title.bold = True
    r_title.font.size = Pt(16)
    r_title.font.color.rgb = RGBColor(15, 23, 42) # Slate 900
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Comprehensive Comparative Breakdown: Features, Performance, Security & Usability Analysis")
    r_sub.font.size = Pt(10.5)
    r_sub.font.color.rgb = RGBColor(71, 85, 105)
    
    doc.add_paragraph() # Spacer
    
    # Executive Summary Paragraph
    p_exec = doc.add_paragraph()
    r_exec_title = p_exec.add_run("Executive Summary:\n")
    r_exec_title.bold = True
    r_exec_title.font.size = Pt(11)
    
    r_exec_text = p_exec.add_run(
        "This engineering report evaluates DevDash (a production-ready native GUI database workspace built with Tauri 2.0, Rust, and React 18) "
        "against DBForge (rishi-jat/dbforge, an experimental local-first database CLI kernel currently at v0.1.0 for SQLite). "
        "Each parameter compares both tools directly and provides a clear technical justification for which is better and why."
    )
    r_exec_text.font.size = Pt(9.5)
    r_exec_text.font.color.rgb = RGBColor(51, 65, 85)
    
    doc.add_paragraph() # Spacer
    
    # 30 Parameters Data: (Parameter, DevDash, DBForge, Advantage, Which Is Better & Why)
    comparison_data = [
        (
            "1. Target Vision & Identity",
            "Full-featured desktop GUI database engineering client & TablePlus alternative.",
            "Local-first database engineering CLI kernel (migration planner & policy checker).",
            "DevDash",
            "DevDash is better because it serves both interactive data exploration/editing and complex schema workflows in a complete desktop application, whereas DBForge is currently restricted to a headless CLI tool."
        ),
        (
            "2. Project Status & Maturity",
            "Production-ready v1.0.0 (25/25 GAPs passed, standalone installers).",
            "Experimental v0.1.0 kernel (headless CLI + SQLite vertical slice).",
            "DevDash",
            "DevDash is better because it is a complete, fully tested product with 25 passed implementation GAPs, ready for daily enterprise use, while DBForge is an early v0.1 experimental prototype."
        ),
        (
            "3. User Interface & Chrome",
            "Rich native GUI (React 18 + Tailwind + Framer Motion + Light/Dark theme).",
            "Headless Command-Line Interface (`dbforge-cli`) — GUI chrome deferred.",
            "DevDash",
            "DevDash is better because its visual GUI features multi-tab workspace orchestration, ERD visualizers, and virtualized data grids that make complex database interaction fast and accessible to all developers."
        ),
        (
            "4. Core Technology Stack",
            "Tauri 2.0 + Rust Engine + React 18 TypeScript Frontend.",
            "Pure Rust Workspace Crates (`dbforge-kernel`, `dbforge-driver-sqlite`).",
            "DevDash",
            "DevDash is better because it combines Rust's memory safety and multi-threading performance with React 18's rich UI ecosystem, giving users both speed and visual quality."
        ),
        (
            "5. Memory & Footprint",
            "Ultra-light ~20MB RAM desktop RAM footprint.",
            "Minimal CLI process memory footprint (<10MB RAM).",
            "Tie",
            "Tie: DBForge is slightly smaller as a headless CLI (<10MB), but DevDash's ~20MB RAM footprint for a complete desktop GUI is extraordinarily light compared to Electron tools (300MB+)."
        ),
        (
            "6. Database Engine Drivers",
            "Postgres, MySQL, MariaDB, SQLite, MSSQL, CockroachDB, Redshift, DuckDB, Oracle, ClickHouse.",
            "SQLite 3 only (`rusqlite` driver in v0.1; Postgres/MySQL deferred).",
            "DevDash",
            "DevDash is better because it natively supports 10+ major SQL database drivers out-of-the-box via `sqlx::AnyPool`, while DBForge only supports single-file SQLite databases in v0.1."
        ),
        (
            "7. NoSQL & Document Support",
            "Native Redis RESP key-value TTL browser & MongoDB BSON collection tree.",
            "Out of scope (Relational SQLite files only).",
            "DevDash",
            "DevDash is better because modern stacks rely heavily on Redis and MongoDB; DevDash provides dedicated NoSQL viewports alongside relational databases."
        ),
        (
            "8. Transactional Edit Staging",
            "Git-style cell diff staging (`old_value → new_value`) with pre-commit review modal.",
            "Non-mutating migration planner (`dbforge plan`) + `dbforge_schema_history` table.",
            "DevDash",
            "DevDash is better because its git-style cell diff staging lets developers inspect exact row modifications before applying them atomically to production databases."
        ),
        (
            "9. Production Write Guards",
            "Safe Mode shield requiring typed `CONFIRM` for destructive queries (`DROP`/`TRUNCATE`).",
            "Requires `--confirm-env prod --reason \"hotfix\"` flag for production writes.",
            "Tie",
            "Tie: Both tools implement excellent production guardrails—DevDash visually intercepts destructive queries with safe mode dialogs, while DBForge requires explicit CLI flags and reasons."
        ),
        (
            "10. Query Streaming & Caps",
            "500-row chunk IPC event streaming (`stream_dynamic_query`) to prevent RAM bloating.",
            "Fetch materialization caps (no automatic SQL query rewriting).",
            "DevDash",
            "DevDash is better because its 500-row chunk IPC streaming allows processing datasets with 100,000+ rows smoothly without freezing the UI or bloating RAM."
        ),
        (
            "11. Visual EXPLAIN Cost Graph",
            "Interactive tree card graph for `EXPLAIN ANALYZE` with scan alerts & buffer hit ratios.",
            "Out of scope.",
            "DevDash",
            "DevDash is better because developers can instantly identify performance bottlenecks, sequential scans, and un-indexed queries through visual cost cards."
        ),
        (
            "12. Observability & Telemetry",
            "6-card Recharts Bento telemetry grid (CPU, RAM, cache hit, locks, slow query log, TCP ping).",
            "Out of scope.",
            "DevDash",
            "DevDash is better because it provides live real-time server health monitoring, active query lock tracking, and slow query diagnostics directly in the workspace."
        ),
        (
            "13. ERD Schema Visualizer",
            "Interactive React Flow force-directed node map with 1-click DDL exporter.",
            "Out of scope.",
            "DevDash",
            "DevDash is better because it auto-generates visual entity-relationship diagrams from database foreign keys with 1-click DDL migration export."
        ),
        (
            "14. Stored Routine Debugger",
            "Parameter inspector & auto-generated `CALL` / `SELECT` SQL executor for procedures.",
            "Out of scope.",
            "DevDash",
            "DevDash is better because it allows inspecting PL/pgSQL, T-SQL, and MySQL stored procedures, functions, and triggers with parameter input forms."
        ),
        (
            "15. User & Privilege Matrix",
            "Visual 7-permission matrix table (`SELECT`, `INSERT`, `UPDATE`, etc.) & `GRANT` SQL generator.",
            "Coarse `read_only` vs write role profiles.",
            "DevDash",
            "DevDash is better because it visualizes complete permission matrices across 7 privilege types and auto-generates exact `GRANT`/`REVOKE` SQL statements."
        ),
        (
            "16. AI & Machine Intelligence",
            "100% offline local Ollama LLM (`qwen2.5-coder`) + Claude + OpenAI + `Cmd+K` bar.",
            "Explicitly out of scope (\"Not product identity\").",
            "DevDash",
            "DevDash is better because its 100% offline local AI integration allows developers to generate SQL from natural language without sending sensitive schema data to the cloud."
        ),
        (
            "17. SOC2 / HIPAA Audit Trail",
            "Append-only JSONL event logger (`audit_log.jsonl`) tracking user, IP, SQL, and rows.",
            "L1 `dbforge_schema_history` table in target DB for applied migrations.",
            "DevDash",
            "DevDash is better because it logs every client query, export, and credentials operation into an append-only JSONL audit trail for SOC2 and HIPAA compliance."
        ),
        (
            "18. Data Masking & PII Protection",
            "Customizable pattern rules (`ssn`, `email`, `credit_card`) with full/partial masking & SHA-256.",
            "Out of scope.",
            "DevDash",
            "DevDash is better because it automatically obfuscates sensitive PII columns locally before rendering, adhering to strict GDPR data protection rules."
        ),
        (
            "19. Live DDL Schema Diff & Sync",
            "Live comparison between `Dev` and `Prod` generating multi-statement migration DDL scripts.",
            "Sequential `.sql` file migration scanner & applier (`dbforge apply`).",
            "DevDash",
            "DevDash is better because it performs live schema comparison between two active databases and outputs clean `ALTER TABLE` DDL sync scripts."
        ),
        (
            "20. Synthetic Data Seed Generator",
            "1-click synthetic seed generator populating 100 to 5,000 realistic rows (names, emails, IPs).",
            "Out of scope.",
            "DevDash",
            "DevDash is better because developers can instantly generate thousands of realistic mock rows matching schema data types to test query performance."
        ),
        (
            "21. Visual No-Code Query Builder",
            "Drag-and-drop block query builder for SELECT, JOIN, WHERE, GROUP BY, and ORDER BY.",
            "Out of scope.",
            "DevDash",
            "DevDash is better because non-SQL team members or quick-query developers can construct complex multi-table joins visually without writing manual SQL."
        ),
        (
            "22. Multi-Format Exporter/Importer",
            "Export to CSV, JSON, SQL dump, JSONL, Markdown Table, and Apache Parquet.",
            "CLI query text output & JSON print.",
            "DevDash",
            "DevDash is better because it supports 6 export formats including high-performance compressed Apache Parquet and streaming JSON Lines."
        ),
        (
            "23. Grid Range Copy/Paste",
            "2D cell block range selection (`Shift+Arrow`) & TSV copy/paste (`Ctrl+C/V`) for Excel.",
            "Out of scope.",
            "DevDash",
            "DevDash is better because it behaves like Excel/Google Sheets, allowing developers to select rectangular blocks of cells and paste TSV data directly into staged edits."
        ),
        (
            "24. Native SSH Tunneling",
            "Thread-safe background TCP listener utilizing native `ssh2` Rust crate.",
            "Out of scope.",
            "DevDash",
            "DevDash is better because it enables secure encrypted connections to remote private databases behind bastion hosts via native SSH port forwarding."
        ),
        (
            "25. Cloud IAM Auth Protocols",
            "Cloud IAM authentication parameter builders for AWS Redshift, GCP BigQuery & Azure SQL.",
            "Out of scope.",
            "DevDash",
            "DevDash is better because it supports modern cloud-native authentication (Service Accounts, STS tokens, OAuth2 AD) required in enterprise cloud deployments."
        ),
        (
            "26. Credentials Encryption",
            "OS Keyring isolation (`keyring` crate) + AES-256-GCM encrypted backup exporter.",
            "Out of scope (Local env/secret files).",
            "DevDash",
            "DevDash is better because connection passwords are isolated in native OS keyrings and backup exports are secured using AES-256-GCM encryption."
        ),
        (
            "27. Developer Shortcuts",
            "Global `Cmd+K` palette, `Cmd+I` formatter, `Cmd+Enter` query run, custom shortcut manager.",
            "CLI command arguments & flags (`--db`, `--role`).",
            "DevDash",
            "DevDash is better because rich keyboard shortcuts dramatically accelerate developer productivity during daily database engineering tasks."
        ),
        (
            "28. Desktop Build Packaging",
            "Standalone installers for Windows (`.msi`/`.exe`), macOS (`.dmg`), and Linux (`.AppImage`/`.deb`).",
            "Cargo build binary executable (`cargo build -p dbforge-cli`).",
            "DevDash",
            "DevDash is better because it packages complete native desktop installers across all major operating systems ready for end-user deployment."
        ),
        (
            "29. OS Security Workaround Guide",
            "Documented Gatekeeper & SmartScreen security bypass guides for unsigned open-source binaries.",
            "Out of scope.",
            "DevDash",
            "DevDash is better because it provides clear, step-by-step documentation for users to bypass unsigned binary warnings on macOS and Windows."
        ),
        (
            "30. Licensing & Governance",
            "Open-source MIT License.",
            "Open-source Apache-2.0 License.",
            "Tie",
            "Tie: Both projects are 100% open-source under permissive software licenses (MIT vs Apache-2.0), ensuring zero vendor lock-in."
        )
    ]

    # Create Table (5 Columns)
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Column Widths (Total ~10 inches on 11-inch landscape page)
    col_widths = [Inches(1.5), Inches(2.2), Inches(1.8), Inches(1.0), Inches(3.5)]

    # Table Header Row
    hdr_cells = table.rows[0].cells
    hdr_titles = ["Parameter", "DevDash (Tauri + Rust GUI)", "DBForge (Rust CLI Kernel)", "Advantage", "Which is Better & Why"]
    
    for i, title in enumerate(hdr_titles):
        cell = hdr_cells[i]
        cell.width = col_widths[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, "0F172A") # Slate 900
        set_cell_margins(cell, top=120, bottom=120, left=100, right=100)

    # Populate Data Rows
    for row_idx, (param, devdash_text, dbforge_text, advantage, why_text) in enumerate(comparison_data):
        row_cells = table.add_row().cells
        bg_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF" # Zebra striping
        
        for i, text in enumerate([param, devdash_text, dbforge_text, advantage, why_text]):
            cell = row_cells[i]
            cell.width = col_widths[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.font.size = Pt(8.5)
            
            if i == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(15, 23, 42)
            elif i == 3:
                run.bold = True
                if "DevDash" in text:
                    run.font.color.rgb = RGBColor(16, 185, 129) # Emerald
                else:
                    run.font.color.rgb = RGBColor(99, 102, 241) # Indigo
            elif i == 4:
                run.font.color.rgb = RGBColor(15, 23, 42)
            else:
                run.font.color.rgb = RGBColor(51, 65, 85)
                
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)

    # Set Borders via XML
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    output_filename = "DevDash_vs_DBForge_Comparison_Report.docx"
    output_path = os.path.join(os.getcwd(), output_filename)
    try:
        if os.path.exists(output_path):
            os.remove(output_path)
    except Exception:
        output_filename = "DevDash_vs_DBForge_5Column_Comparison_Report.docx"
        output_path = os.path.join(os.getcwd(), output_filename)

    doc.save(output_path)
    print(f"Successfully generated 5-column comparison report document: {output_path}")

if __name__ == "__main__":
    create_comparison_docx()
